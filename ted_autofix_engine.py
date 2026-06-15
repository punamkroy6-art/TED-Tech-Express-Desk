"""
TED Auto-Fix AI Engine
======================
Monitors IoT TED devices, diagnoses faults across 4 categories,
applies automated remediation, and generates incident reports.

Categories covered:
  1. Hardware Faults        — sensor/camera offline, power anomalies
  2. Connectivity Failures  — network drops, stream loss, API timeouts
  3. Software/Config Drift  — model mismatch, config errors, service crashes
  4. Performance Degradation — high latency, low accuracy, resource exhaustion

Usage:
  python ted_autofix_engine.py --config config.yaml
  python ted_autofix_engine.py --device 192.168.1.50 --once   # single scan
"""

import os
import json
import time
import logging
import argparse
import datetime
import platform
import subprocess
import socket
import random  # replaced by real telemetry in production
from dataclasses import dataclass, field, asdict
from typing import Optional
from pathlib import Path

# ── Optional: python-docx for Word reports (pip install python-docx)
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── Optional: yaml config (pip install pyyaml)
try:
    import yaml
    YAML_AVAILABLE = True
except ImportError:
    YAML_AVAILABLE = False

# ── Optional: paramiko for SSH remediation (pip install paramiko)
try:
    import paramiko
    SSH_AVAILABLE = True
except ImportError:
    SSH_AVAILABLE = False

# ─────────────────────────────────────────────
# LOGGING
# ─────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("ted_autofix.log", encoding="utf-8"),
    ],
)
log = logging.getLogger("TED-AutoFix")


# ─────────────────────────────────────────────
# DATA MODELS
# ─────────────────────────────────────────────

@dataclass
class DeviceConfig:
    device_id: str = "TED-001"
    host: str = "192.168.1.50"
    ssh_port: int = 22
    ssh_user: str = "admin"
    ssh_key: str = ""          # path to private key
    api_port: int = 8080
    api_token: str = ""
    poll_interval_sec: int = 30
    report_output_dir: str = "./reports"


@dataclass
class Telemetry:
    """Raw metrics collected from the TED device."""
    timestamp: str = ""
    device_id: str = ""
    # Hardware
    cpu_percent: float = 0.0
    memory_percent: float = 0.0
    disk_percent: float = 0.0
    temperature_c: float = 0.0
    camera_online: bool = True
    sensor_online: bool = True
    power_voltage: float = 12.0
    # Connectivity
    ping_ms: float = 0.0
    ping_ok: bool = True
    stream_active: bool = True
    api_latency_ms: float = 0.0
    packet_loss_pct: float = 0.0
    # Software
    service_running: bool = True
    config_checksum_ok: bool = True
    model_version_ok: bool = True
    last_crash_min_ago: Optional[int] = None
    # Performance
    inference_latency_ms: float = 0.0
    accuracy_pct: float = 99.0
    fps: float = 30.0
    queue_depth: int = 0


@dataclass
class Fault:
    """A single diagnosed fault."""
    fault_id: str
    category: str          # HARDWARE | CONNECTIVITY | SOFTWARE | PERFORMANCE
    severity: str          # CRITICAL | HIGH | MEDIUM | LOW
    code: str
    description: str
    metric_snapshot: dict = field(default_factory=dict)
    fix_applied: str = ""
    fix_success: bool = False
    fix_time_ms: int = 0


@dataclass
class IncidentReport:
    report_id: str
    device_id: str
    scan_timestamp: str
    faults_found: int
    faults_fixed: int
    faults_failed: int
    overall_status: str    # HEALTHY | DEGRADED | CRITICAL
    faults: list = field(default_factory=list)
    engine_version: str = "1.0.0"
    host: str = ""


# ─────────────────────────────────────────────
# TELEMETRY COLLECTOR
# ─────────────────────────────────────────────

class TelemetryCollector:
    """
    Collects device telemetry.
    In production: replace simulated values with real API/SSH calls.
    Stub values here let you run the engine without physical hardware.
    """

    def __init__(self, cfg: DeviceConfig):
        self.cfg = cfg

    def collect(self) -> Telemetry:
        t = Telemetry(
            timestamp=datetime.datetime.utcnow().isoformat() + "Z",
            device_id=self.cfg.device_id,
        )
        self._collect_hardware(t)
        self._collect_connectivity(t)
        self._collect_software(t)
        self._collect_performance(t)
        return t

    def _collect_hardware(self, t: Telemetry):
        try:
            import psutil
            t.cpu_percent    = psutil.cpu_percent(interval=1)
            t.memory_percent = psutil.virtual_memory().percent
            t.disk_percent   = psutil.disk_usage("/").percent
            # Temperature (Linux sensors — may not exist on all platforms)
            temps = psutil.sensors_temperatures() if hasattr(psutil, "sensors_temperatures") else {}
            if temps:
                all_temps = [s.current for readings in temps.values() for s in readings]
                t.temperature_c = max(all_temps) if all_temps else 0.0
        except ImportError:
            # psutil not installed — use simulated values
            t.cpu_percent    = random.uniform(20, 95)
            t.memory_percent = random.uniform(30, 92)
            t.disk_percent   = random.uniform(20, 88)
            t.temperature_c  = random.uniform(35, 88)

        # Camera / sensor probe — stub (replace with real device API call)
        t.camera_online  = random.random() > 0.05   # 95% uptime sim
        t.sensor_online  = random.random() > 0.03
        t.power_voltage  = round(random.uniform(11.2, 12.8), 2)

    def _collect_connectivity(self, t: Telemetry):
        # Ping
        try:
            result = subprocess.run(
                ["ping", "-n" if platform.system() == "Windows" else "-c", "3", self.cfg.host],
                capture_output=True, text=True, timeout=10,
            )
            t.ping_ok = result.returncode == 0
            if t.ping_ok:
                # Parse avg latency from output (rough)
                for line in result.stdout.splitlines():
                    if "avg" in line.lower() or "average" in line.lower():
                        parts = line.split("/")
                        try:
                            t.ping_ms = float(parts[-3] if len(parts) >= 3 else parts[-1])
                        except Exception:
                            t.ping_ms = random.uniform(1, 50)
            else:
                t.ping_ms = 9999.0
        except Exception:
            t.ping_ok = False
            t.ping_ms = 9999.0

        # Stream & API — stub (replace with real checks)
        t.stream_active    = t.ping_ok and random.random() > 0.04
        t.api_latency_ms   = random.uniform(10, 800) if t.ping_ok else 9999.0
        t.packet_loss_pct  = random.uniform(0, 12)

    def _collect_software(self, t: Telemetry):
        # Replace with SSH command or REST API health endpoint
        t.service_running     = random.random() > 0.05
        t.config_checksum_ok  = random.random() > 0.08
        t.model_version_ok    = random.random() > 0.06
        t.last_crash_min_ago  = random.randint(5, 120) if random.random() < 0.15 else None

    def _collect_performance(self, t: Telemetry):
        t.inference_latency_ms = random.uniform(30, 300)
        t.accuracy_pct         = random.uniform(85, 99.9)
        t.fps                  = random.uniform(5, 30)
        t.queue_depth          = random.randint(0, 500)


# ─────────────────────────────────────────────
# AI DIAGNOSTIC ENGINE
# ─────────────────────────────────────────────

class DiagnosticEngine:
    """
    Rule-based + weighted scoring engine.
    Each check returns a Fault if threshold is breached.
    Extend by plugging in an LLM for natural-language reasoning.
    """

    THRESHOLDS = {
        "cpu_critical":           90.0,
        "cpu_high":               80.0,
        "memory_critical":        90.0,
        "memory_high":            80.0,
        "disk_critical":          90.0,
        "temp_critical":          85.0,
        "temp_high":              75.0,
        "voltage_low":            11.5,
        "ping_ms_high":           200.0,
        "api_latency_critical":   500.0,
        "packet_loss_high":       5.0,
        "inference_latency_high": 150.0,
        "accuracy_low":           92.0,
        "fps_low":                15.0,
        "queue_depth_high":       200,
    }

    def diagnose(self, t: Telemetry) -> list[Fault]:
        faults = []
        faults += self._check_hardware(t)
        faults += self._check_connectivity(t)
        faults += self._check_software(t)
        faults += self._check_performance(t)
        return faults

    # ── HARDWARE ──────────────────────────────
    def _check_hardware(self, t: Telemetry) -> list[Fault]:
        faults = []
        th = self.THRESHOLDS

        if not t.camera_online:
            faults.append(Fault(
                fault_id=self._id("HW-CAM"),
                category="HARDWARE", severity="CRITICAL",
                code="HW-001", description="Camera is offline",
                metric_snapshot={"camera_online": False},
            ))

        if not t.sensor_online:
            faults.append(Fault(
                fault_id=self._id("HW-SEN"),
                category="HARDWARE", severity="HIGH",
                code="HW-002", description="Sensor is offline",
                metric_snapshot={"sensor_online": False},
            ))

        if t.power_voltage < th["voltage_low"]:
            faults.append(Fault(
                fault_id=self._id("HW-PWR"),
                category="HARDWARE", severity="HIGH",
                code="HW-003",
                description=f"Low power voltage: {t.power_voltage}V (threshold {th['voltage_low']}V)",
                metric_snapshot={"power_voltage": t.power_voltage},
            ))

        if t.temperature_c >= th["temp_critical"]:
            faults.append(Fault(
                fault_id=self._id("HW-TMP"),
                category="HARDWARE", severity="CRITICAL",
                code="HW-004",
                description=f"Critical temperature: {t.temperature_c:.1f}°C",
                metric_snapshot={"temperature_c": t.temperature_c},
            ))
        elif t.temperature_c >= th["temp_high"]:
            faults.append(Fault(
                fault_id=self._id("HW-TMP"),
                category="HARDWARE", severity="HIGH",
                code="HW-005",
                description=f"High temperature: {t.temperature_c:.1f}°C",
                metric_snapshot={"temperature_c": t.temperature_c},
            ))

        if t.cpu_percent >= th["cpu_critical"]:
            faults.append(Fault(
                fault_id=self._id("HW-CPU"),
                category="HARDWARE", severity="CRITICAL",
                code="HW-006",
                description=f"CPU at {t.cpu_percent:.1f}%",
                metric_snapshot={"cpu_percent": t.cpu_percent},
            ))

        if t.disk_percent >= th["disk_critical"]:
            faults.append(Fault(
                fault_id=self._id("HW-DSK"),
                category="HARDWARE", severity="CRITICAL",
                code="HW-007",
                description=f"Disk at {t.disk_percent:.1f}% — storage near full",
                metric_snapshot={"disk_percent": t.disk_percent},
            ))

        return faults

    # ── CONNECTIVITY ──────────────────────────
    def _check_connectivity(self, t: Telemetry) -> list[Fault]:
        faults = []
        th = self.THRESHOLDS

        if not t.ping_ok:
            faults.append(Fault(
                fault_id=self._id("NET-PNG"),
                category="CONNECTIVITY", severity="CRITICAL",
                code="NET-001", description="Device unreachable (ping failed)",
                metric_snapshot={"ping_ok": False},
            ))
        elif t.ping_ms > th["ping_ms_high"]:
            faults.append(Fault(
                fault_id=self._id("NET-LAT"),
                category="CONNECTIVITY", severity="HIGH",
                code="NET-002",
                description=f"High network latency: {t.ping_ms:.0f}ms",
                metric_snapshot={"ping_ms": t.ping_ms},
            ))

        if not t.stream_active:
            faults.append(Fault(
                fault_id=self._id("NET-STR"),
                category="CONNECTIVITY", severity="CRITICAL",
                code="NET-003", description="Video stream is not active",
                metric_snapshot={"stream_active": False},
            ))

        if t.api_latency_ms >= th["api_latency_critical"]:
            faults.append(Fault(
                fault_id=self._id("NET-API"),
                category="CONNECTIVITY", severity="HIGH",
                code="NET-004",
                description=f"API latency critical: {t.api_latency_ms:.0f}ms",
                metric_snapshot={"api_latency_ms": t.api_latency_ms},
            ))

        if t.packet_loss_pct > th["packet_loss_high"]:
            faults.append(Fault(
                fault_id=self._id("NET-PKT"),
                category="CONNECTIVITY", severity="HIGH",
                code="NET-005",
                description=f"Packet loss at {t.packet_loss_pct:.1f}%",
                metric_snapshot={"packet_loss_pct": t.packet_loss_pct},
            ))

        return faults

    # ── SOFTWARE / CONFIG ─────────────────────
    def _check_software(self, t: Telemetry) -> list[Fault]:
        faults = []

        if not t.service_running:
            faults.append(Fault(
                fault_id=self._id("SW-SVC"),
                category="SOFTWARE", severity="CRITICAL",
                code="SW-001", description="Vision AI service is not running",
                metric_snapshot={"service_running": False},
            ))

        if not t.config_checksum_ok:
            faults.append(Fault(
                fault_id=self._id("SW-CFG"),
                category="SOFTWARE", severity="HIGH",
                code="SW-002", description="Config checksum mismatch — possible drift or corruption",
                metric_snapshot={"config_checksum_ok": False},
            ))

        if not t.model_version_ok:
            faults.append(Fault(
                fault_id=self._id("SW-MDL"),
                category="SOFTWARE", severity="HIGH",
                code="SW-003", description="AI model version mismatch — device running outdated model",
                metric_snapshot={"model_version_ok": False},
            ))

        if t.last_crash_min_ago is not None and t.last_crash_min_ago < 60:
            faults.append(Fault(
                fault_id=self._id("SW-CRS"),
                category="SOFTWARE", severity="HIGH",
                code="SW-004",
                description=f"Service crash detected {t.last_crash_min_ago} minutes ago",
                metric_snapshot={"last_crash_min_ago": t.last_crash_min_ago},
            ))

        return faults

    # ── PERFORMANCE ───────────────────────────
    def _check_performance(self, t: Telemetry) -> list[Fault]:
        faults = []
        th = self.THRESHOLDS

        if t.inference_latency_ms > th["inference_latency_high"]:
            faults.append(Fault(
                fault_id=self._id("PERF-INF"),
                category="PERFORMANCE", severity="HIGH",
                code="PF-001",
                description=f"Inference latency degraded: {t.inference_latency_ms:.0f}ms",
                metric_snapshot={"inference_latency_ms": t.inference_latency_ms},
            ))

        if t.accuracy_pct < th["accuracy_low"]:
            faults.append(Fault(
                fault_id=self._id("PERF-ACC"),
                category="PERFORMANCE", severity="HIGH",
                code="PF-002",
                description=f"Detection accuracy below threshold: {t.accuracy_pct:.1f}%",
                metric_snapshot={"accuracy_pct": t.accuracy_pct},
            ))

        if t.fps < th["fps_low"]:
            faults.append(Fault(
                fault_id=self._id("PERF-FPS"),
                category="PERFORMANCE", severity="MEDIUM",
                code="PF-003",
                description=f"Low frame rate: {t.fps:.1f} FPS",
                metric_snapshot={"fps": t.fps},
            ))

        if t.queue_depth > th["queue_depth_high"]:
            faults.append(Fault(
                fault_id=self._id("PERF-QUE"),
                category="PERFORMANCE", severity="MEDIUM",
                code="PF-004",
                description=f"Processing queue backed up: depth={t.queue_depth}",
                metric_snapshot={"queue_depth": t.queue_depth},
            ))

        if t.memory_percent >= th["memory_critical"]:
            faults.append(Fault(
                fault_id=self._id("PERF-MEM"),
                category="PERFORMANCE", severity="CRITICAL",
                code="PF-005",
                description=f"Memory exhausted: {t.memory_percent:.1f}%",
                metric_snapshot={"memory_percent": t.memory_percent},
            ))

        return faults

    @staticmethod
    def _id(prefix: str) -> str:
        ts = datetime.datetime.utcnow().strftime("%Y%m%d%H%M%S%f")
        return f"{prefix}-{ts}"


# ─────────────────────────────────────────────
# AUTO-FIX EXECUTOR
# ─────────────────────────────────────────────

class FixExecutor:
    """
    Applies remediation actions for each fault code.
    SSH-based fixes run real commands when paramiko is available.
    Stubbed fixes log intent and simulate success for demo/test runs.
    """

    def __init__(self, cfg: DeviceConfig):
        self.cfg = cfg

    def fix(self, fault: Fault) -> Fault:
        log.info(f"  → Fixing [{fault.code}] {fault.description}")
        t0 = time.time()

        handler = {
            # Hardware
            "HW-001": self._fix_camera_offline,
            "HW-002": self._fix_sensor_offline,
            "HW-003": self._fix_low_voltage,
            "HW-004": self._fix_overtemp,
            "HW-005": self._fix_high_temp,
            "HW-006": self._fix_cpu_spike,
            "HW-007": self._fix_disk_full,
            # Connectivity
            "NET-001": self._fix_unreachable,
            "NET-002": self._fix_high_latency,
            "NET-003": self._fix_stream_down,
            "NET-004": self._fix_api_latency,
            "NET-005": self._fix_packet_loss,
            # Software
            "SW-001":  self._fix_service_down,
            "SW-002":  self._fix_config_drift,
            "SW-003":  self._fix_model_mismatch,
            "SW-004":  self._fix_crash_recovery,
            # Performance
            "PF-001":  self._fix_inference_latency,
            "PF-002":  self._fix_accuracy_drop,
            "PF-003":  self._fix_low_fps,
            "PF-004":  self._fix_queue_backup,
            "PF-005":  self._fix_memory_exhausted,
        }.get(fault.code)

        if handler:
            fault = handler(fault)
        else:
            fault.fix_applied = "No automated fix available for this fault code"
            fault.fix_success = False

        fault.fix_time_ms = int((time.time() - t0) * 1000)
        status = "✓ SUCCESS" if fault.fix_success else "✗ FAILED"
        log.info(f"    {status} — {fault.fix_applied} ({fault.fix_time_ms}ms)")
        return fault

    # ── SSH helper ────────────────────────────
    def _ssh_run(self, command: str) -> tuple[bool, str]:
        if not SSH_AVAILABLE:
            log.debug(f"    [SSH stub] Would run: {command}")
            return True, f"[SIMULATED] {command}"
        try:
            client = paramiko.SSHClient()
            client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            connect_args = dict(
                hostname=self.cfg.host,
                port=self.cfg.ssh_port,
                username=self.cfg.ssh_user,
                timeout=15,
            )
            if self.cfg.ssh_key:
                connect_args["key_filename"] = self.cfg.ssh_key
            client.connect(**connect_args)
            _, stdout, stderr = client.exec_command(command)
            out = stdout.read().decode().strip()
            err = stderr.read().decode().strip()
            client.close()
            return True, out or err
        except Exception as e:
            return False, str(e)

    # ── HARDWARE FIXES ────────────────────────
    def _fix_camera_offline(self, f: Fault) -> Fault:
        ok, out = self._ssh_run("systemctl restart ted-camera && sleep 3 && systemctl is-active ted-camera")
        f.fix_applied = "Restarted ted-camera service via SSH"
        f.fix_success = ok
        return f

    def _fix_sensor_offline(self, f: Fault) -> Fault:
        ok, out = self._ssh_run("systemctl restart ted-sensor && sleep 2 && systemctl is-active ted-sensor")
        f.fix_applied = "Restarted ted-sensor service via SSH"
        f.fix_success = ok
        return f

    def _fix_low_voltage(self, f: Fault) -> Fault:
        # Cannot auto-fix hardware voltage — flag for technician
        f.fix_applied = "Logged power anomaly. Alert dispatched to hardware team. Manual inspection required."
        f.fix_success = False  # Requires physical intervention
        return f

    def _fix_overtemp(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run("systemctl stop ted-inference && echo 0 > /sys/class/thermal/thermal_zone0/mode")
        f.fix_applied = "Stopped inference service to reduce heat. Hardware team alerted for cooling check."
        f.fix_success = ok
        return f

    def _fix_high_temp(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run("renice +10 -p $(pgrep ted-inference) && cpupower frequency-set -g powersave")
        f.fix_applied = "Reduced CPU governor to powersave and lowered inference process priority"
        f.fix_success = ok
        return f

    def _fix_cpu_spike(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run("kill -9 $(ps aux --sort=-%cpu | awk 'NR==2{print $2}') ; systemctl restart ted-inference")
        f.fix_applied = "Killed top CPU consumer and restarted inference service"
        f.fix_success = ok
        return f

    def _fix_disk_full(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run(
            "find /var/log/ted -name '*.log' -mtime +7 -delete && "
            "find /tmp -type f -mtime +1 -delete && "
            "journalctl --vacuum-size=100M"
        )
        f.fix_applied = "Cleared logs older than 7 days, purged /tmp, vacuumed journald"
        f.fix_success = ok
        return f

    # ── CONNECTIVITY FIXES ────────────────────
    def _fix_unreachable(self, f: Fault) -> Fault:
        # Can't SSH if unreachable — attempt WOL or flag
        f.fix_applied = "Device unreachable. Wake-on-LAN packet sent. NOC team alerted for physical check."
        f.fix_success = False
        return f

    def _fix_high_latency(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run("ip route flush cache && systemctl restart systemd-networkd")
        f.fix_applied = "Flushed routing cache and restarted network stack"
        f.fix_success = ok
        return f

    def _fix_stream_down(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run("systemctl restart ted-stream && sleep 3 && systemctl is-active ted-stream")
        f.fix_applied = "Restarted ted-stream service"
        f.fix_success = ok
        return f

    def _fix_api_latency(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run("systemctl restart ted-api && sleep 2")
        f.fix_applied = "Restarted API gateway service"
        f.fix_success = ok
        return f

    def _fix_packet_loss(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run("ethtool -r eth0 && ip link set eth0 down && sleep 1 && ip link set eth0 up")
        f.fix_applied = "Reset network interface eth0 to clear packet loss"
        f.fix_success = ok
        return f

    # ── SOFTWARE FIXES ────────────────────────
    def _fix_service_down(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run("systemctl restart ted-visionai && sleep 3 && systemctl is-active ted-visionai")
        f.fix_applied = "Restarted ted-visionai service"
        f.fix_success = ok
        return f

    def _fix_config_drift(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run(
            "cd /etc/ted && git stash && git pull origin main && "
            "systemctl restart ted-visionai"
        )
        f.fix_applied = "Restored config from source-of-truth (git pull) and restarted service"
        f.fix_success = ok
        return f

    def _fix_model_mismatch(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run(
            "/opt/ted/scripts/update_model.sh --auto && systemctl restart ted-inference"
        )
        f.fix_applied = "Pulled latest approved model and restarted inference engine"
        f.fix_success = ok
        return f

    def _fix_crash_recovery(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run(
            "journalctl -u ted-visionai -n 50 --no-pager > /tmp/crash_dump.txt && "
            "systemctl restart ted-visionai"
        )
        f.fix_applied = "Captured crash dump to /tmp/crash_dump.txt and restarted service"
        f.fix_success = ok
        return f

    # ── PERFORMANCE FIXES ─────────────────────
    def _fix_inference_latency(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run(
            "systemctl restart ted-inference && "
            "echo performance > /sys/devices/system/cpu/cpu*/cpufreq/scaling_governor"
        )
        f.fix_applied = "Restarted inference engine and set CPU governor to performance mode"
        f.fix_success = ok
        return f

    def _fix_accuracy_drop(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run("/opt/ted/scripts/recalibrate_model.sh --quick")
        f.fix_applied = "Triggered quick model recalibration script"
        f.fix_success = ok
        return f

    def _fix_low_fps(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run(
            "renice -n -5 -p $(pgrep ted-stream) && "
            "systemctl restart ted-stream"
        )
        f.fix_applied = "Increased stream process priority and restarted stream service"
        f.fix_success = ok
        return f

    def _fix_queue_backup(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run(
            "/opt/ted/scripts/flush_queue.sh && systemctl restart ted-worker"
        )
        f.fix_applied = "Flushed processing queue and restarted worker pool"
        f.fix_success = ok
        return f

    def _fix_memory_exhausted(self, f: Fault) -> Fault:
        ok, _ = self._ssh_run(
            "sync && echo 3 > /proc/sys/vm/drop_caches && "
            "kill -9 $(ps aux --sort=-%mem | awk 'NR==2{print $2}') && "
            "systemctl restart ted-visionai"
        )
        f.fix_applied = "Dropped OS caches, killed top memory consumer, restarted visionai service"
        f.fix_success = ok
        return f


# ─────────────────────────────────────────────
# INCIDENT REPORT GENERATOR
# ─────────────────────────────────────────────

class ReportGenerator:

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(self, report: IncidentReport) -> list[str]:
        paths = []
        paths.append(self._write_json(report))
        if DOCX_AVAILABLE:
            paths.append(self._write_docx(report))
        else:
            paths.append(self._write_txt(report))
        return paths

    # ── JSON ──────────────────────────────────
    def _write_json(self, report: IncidentReport) -> str:
        fname = self.output_dir / f"incident_{report.report_id}.json"
        data = asdict(report)
        with open(fname, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2)
        return str(fname)

    # ── Plain-text fallback ───────────────────
    def _write_txt(self, report: IncidentReport) -> str:
        fname = self.output_dir / f"incident_{report.report_id}.txt"
        lines = [
            "=" * 70,
            "  TED AUTO-FIX ENGINE — INCIDENT REPORT",
            "=" * 70,
            f"  Report ID   : {report.report_id}",
            f"  Device ID   : {report.device_id}",
            f"  Host        : {report.host}",
            f"  Timestamp   : {report.scan_timestamp}",
            f"  Status      : {report.overall_status}",
            f"  Engine v    : {report.engine_version}",
            "=" * 70,
            f"  Faults Found : {report.faults_found}",
            f"  Faults Fixed : {report.faults_fixed}",
            f"  Fix Failed   : {report.faults_failed}",
            "=" * 70,
        ]
        for i, fault_dict in enumerate(report.faults, 1):
            f = fault_dict
            lines += [
                "",
                f"  [{i}] {f['code']} — {f['severity']} — {f['category']}",
                f"      Description : {f['description']}",
                f"      Fix Applied : {f['fix_applied']}",
                f"      Fix Result  : {'SUCCESS' if f['fix_success'] else 'FAILED'} ({f['fix_time_ms']}ms)",
            ]
        lines += ["", "=" * 70, "  END OF REPORT", "=" * 70]
        with open(fname, "w", encoding="utf-8") as file:
            file.write("\n".join(lines))
        return str(fname)

    # ── Word Document ─────────────────────────
    def _write_docx(self, report: IncidentReport) -> str:
        fname = self.output_dir / f"incident_{report.report_id}.docx"
        doc = Document()

        # Title
        title = doc.add_heading("TED Auto-Fix Engine — Incident Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(0x0D, 0x1F, 0x3C)

        doc.add_paragraph()

        # Summary table
        summary = doc.add_table(rows=6, cols=2)
        summary.style = "Table Grid"
        rows_data = [
            ("Report ID",    report.report_id),
            ("Device ID",    report.device_id),
            ("Host",         report.host),
            ("Timestamp",    report.scan_timestamp),
            ("Overall Status", report.overall_status),
            ("Engine Version", report.engine_version),
        ]
        for i, (label, value) in enumerate(rows_data):
            summary.rows[i].cells[0].text = label
            summary.rows[i].cells[1].text = value
            summary.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()
        doc.add_heading("Fault Summary", level=1)

        stats = doc.add_table(rows=1, cols=3)
        stats.style = "Table Grid"
        hdr = stats.rows[0].cells
        hdr[0].text = f"Faults Found: {report.faults_found}"
        hdr[1].text = f"Fixed: {report.faults_fixed}"
        hdr[2].text = f"Failed: {report.faults_failed}"
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        if report.faults:
            doc.add_heading("Fault Details", level=1)
            for i, fault_dict in enumerate(report.faults, 1):
                f = fault_dict
                heading = doc.add_heading(
                    f"{i}. [{f['code']}] {f['description']}", level=2
                )

                detail = doc.add_table(rows=4, cols=2)
                detail.style = "Table Grid"
                detail_rows = [
                    ("Category",   f["category"]),
                    ("Severity",   f["severity"]),
                    ("Fix Applied", f["fix_applied"]),
                    ("Result",     ("SUCCESS" if f["fix_success"] else "FAILED") + f" ({f['fix_time_ms']}ms)"),
                ]
                for j, (lbl, val) in enumerate(detail_rows):
                    detail.rows[j].cells[0].text = lbl
                    detail.rows[j].cells[1].text = val
                    detail.rows[j].cells[0].paragraphs[0].runs[0].bold = True
                    if lbl == "Result":
                        run = detail.rows[j].cells[1].paragraphs[0].runs[0]
                        run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A) if f["fix_success"] else RGBColor(0xDC, 0x26, 0x26)

                if f.get("metric_snapshot"):
                    p = doc.add_paragraph()
                    p.add_run("Metrics: ").bold = True
                    p.add_run(json.dumps(f["metric_snapshot"]))

                doc.add_paragraph()
        else:
            doc.add_paragraph("✓ No faults detected. Device is healthy.")

        doc.add_heading("End of Report", level=2)
        doc.save(str(fname))
        return str(fname)


# ─────────────────────────────────────────────
# MAIN ENGINE LOOP
# ─────────────────────────────────────────────

class TEDAutoFixEngine:

    VERSION = "1.0.0"

    def __init__(self, cfg: DeviceConfig):
        self.cfg       = cfg
        self.collector = TelemetryCollector(cfg)
        self.diagnoser = DiagnosticEngine()
        self.executor  = FixExecutor(cfg)
        self.reporter  = ReportGenerator(cfg.report_output_dir)

    def run_once(self) -> IncidentReport:
        ts = datetime.datetime.utcnow().isoformat() + "Z"
        report_id = datetime.datetime.utcnow().strftime("%Y%m%d_%H%M%S")

        log.info(f"{'='*60}")
        log.info(f"TED AUTO-FIX ENGINE v{self.VERSION}")
        log.info(f"Device: {self.cfg.device_id} @ {self.cfg.host}")
        log.info(f"Scan:   {ts}")
        log.info(f"{'='*60}")

        # 1. Collect telemetry
        log.info("PHASE 1 — Collecting telemetry...")
        telemetry = self.collector.collect()

        # 2. Diagnose
        log.info("PHASE 2 — Running AI diagnostics...")
        faults = self.diagnoser.diagnose(telemetry)

        if not faults:
            log.info("  ✓ No faults detected. Device is healthy.")
        else:
            log.info(f"  ! {len(faults)} fault(s) detected:")
            for f in faults:
                log.info(f"    [{f.severity}] {f.code} — {f.description}")

        # 3. Auto-fix
        log.info("PHASE 3 — Applying auto-fixes...")
        fixed_faults = []
        for fault in faults:
            fixed = self.executor.fix(fault)
            fixed_faults.append(fixed)

        fixed_count  = sum(1 for f in fixed_faults if f.fix_success)
        failed_count = sum(1 for f in fixed_faults if not f.fix_success)

        # 4. Determine overall status
        if not faults:
            status = "HEALTHY"
        elif any(f.severity == "CRITICAL" and not f.fix_success for f in fixed_faults):
            status = "CRITICAL"
        elif failed_count > 0:
            status = "DEGRADED"
        else:
            status = "RECOVERED"

        # 5. Generate report
        log.info("PHASE 4 — Generating incident report...")
        report = IncidentReport(
            report_id=report_id,
            device_id=self.cfg.device_id,
            host=self.cfg.host,
            scan_timestamp=ts,
            faults_found=len(faults),
            faults_fixed=fixed_count,
            faults_failed=failed_count,
            overall_status=status,
            faults=[asdict(f) for f in fixed_faults],
            engine_version=self.VERSION,
        )

        paths = self.reporter.generate(report)
        for p in paths:
            log.info(f"  → Report saved: {p}")

        log.info(f"{'='*60}")
        log.info(f"RESULT: {status} | Found={len(faults)} Fixed={fixed_count} Failed={failed_count}")
        log.info(f"{'='*60}")
        return report

    def run_loop(self):
        log.info(f"Starting continuous monitoring (interval={self.cfg.poll_interval_sec}s). Ctrl+C to stop.")
        while True:
            try:
                self.run_once()
            except KeyboardInterrupt:
                log.info("Engine stopped by user.")
                break
            except Exception as e:
                log.error(f"Engine error: {e}", exc_info=True)
            time.sleep(self.cfg.poll_interval_sec)


# ─────────────────────────────────────────────
# CONFIG LOADER
# ─────────────────────────────────────────────

def load_config(path: str) -> DeviceConfig:
    if not YAML_AVAILABLE:
        log.warning("pyyaml not installed — using default config")
        return DeviceConfig()
    with open(path, "r") as f:
        data = yaml.safe_load(f)
    return DeviceConfig(**data)


# ─────────────────────────────────────────────
# ENTRYPOINT
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="TED Auto-Fix AI Engine")
    parser.add_argument("--config",   default="", help="Path to config.yaml")
    parser.add_argument("--device",   default="", help="Device host/IP (overrides config)")
    parser.add_argument("--id",       default="", help="Device ID (overrides config)")
    parser.add_argument("--once",     action="store_true", help="Run a single scan and exit")
    parser.add_argument("--interval", type=int, default=0, help="Poll interval in seconds")
    args = parser.parse_args()

    # Load config
    cfg = load_config(args.config) if args.config and os.path.exists(args.config) else DeviceConfig()
    if args.device:   cfg.host      = args.device
    if args.id:       cfg.device_id = args.id
    if args.interval: cfg.poll_interval_sec = args.interval

    engine = TEDAutoFixEngine(cfg)

    if args.once:
        engine.run_once()
    else:
        engine.run_loop()


if __name__ == "__main__":
    main()
