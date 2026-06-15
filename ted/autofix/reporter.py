"""
reporter.py — Incident report generation via python-docx
Produces a Word .docx document summarising the TED kiosk session,
diagnostics, issue found, fix applied, and outcome.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette matching TED's dark-navy brand ──────────────────────────
NAVY    = RGBColor(0x0D, 0x1B, 0x2A)
BLUE    = RGBColor(0x2E, 0x75, 0xB6)
CYAN    = RGBColor(0x22, 0xD3, 0xEE)
GREEN   = RGBColor(0x22, 0xC5, 0x5E)
RED     = RGBColor(0xEF, 0x44, 0x44)
GREY    = RGBColor(0x64, 0x74, 0x8B)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
BLACK   = RGBColor(0x0F, 0x17, 0x2A)


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.color.rgb = BLUE if level == 1 else NAVY
    return p


def _kv(doc: Document, key: str, value: str, value_color: RGBColor = None):
    """Key: Value paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r_key = p.add_run(f"{key}: ")
    r_key.bold = True
    r_key.font.color.rgb = GREY
    r_key.font.size = Pt(10)
    r_val = p.add_run(str(value))
    r_val.font.color.rgb = value_color or BLACK
    r_val.font.size = Pt(10)


def generate_report(
    session_id: str,
    employee: dict,
    metrics: dict,
    triggered_rules: list,
    fix_results: list,
    diagnosis: dict,
    outcome: str,
    output_dir: str = "reports",
) -> str:
    """
    Generate a Word incident report for a TED kiosk session.
    Returns the path to the saved .docx file.
    """
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Header banner ────────────────────────────────────────────────────
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hdr.add_run("💻  TED — Tech Express Desk")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("IT Support Kiosk — Incident Report")
    rs.font.size = Pt(11)
    rs.font.color.rgb = GREY
    rs.italic = True

    doc.add_paragraph()  # spacer

    # ── Session info ─────────────────────────────────────────────────────
    _heading(doc, "Session Information")
    _kv(doc, "Session ID",  session_id)
    _kv(doc, "Generated",   datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"))
    _kv(doc, "Employee",    employee.get("name", "Unknown"))
    _kv(doc, "Employee ID", employee.get("id", "—"))
    _kv(doc, "Department",  employee.get("dept", "—"))
    _kv(doc, "Outcome",     outcome.upper(),
        value_color=GREEN if outcome == "resolved" else (RED if outcome == "ticket_created" else GREY))

    # ── Hardware diagnostics ─────────────────────────────────────────────
    _heading(doc, "Hardware Diagnostics")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for cell, label in zip(hdr_cells, ["Metric", "Value", "Status"]):
        _set_cell_bg(cell, "0D1B2A")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)

    cpu   = metrics.get("cpu", {})
    mem   = metrics.get("memory", {})
    disk  = metrics.get("disk", {})
    batt  = metrics.get("battery", {})

    rows_data = [
        ("CPU Usage",    f"{cpu.get('percent', '?')}%",       "⚠" if cpu.get("percent", 0) > 85 else "✓"),
        ("RAM Usage",    f"{mem.get('percent', '?')}%  ({mem.get('used_gb','?')}/{mem.get('total_gb','?')} GB)", "⚠" if mem.get("percent", 0) > 88 else "✓"),
        ("Disk Usage",   f"{disk.get('percent', '?')}%  ({disk.get('free_gb','?')} GB free)", "⚠" if disk.get("free_gb", 99) < 5 else "✓"),
        ("Battery",      f"{batt.get('percent', 'N/A')}%" if batt.get("percent") else "Desktop / N/A", "✓"),
        ("Hostname",     metrics.get("hostname", "—"),         "—"),
        ("Platform",     metrics.get("platform", "—"),         "—"),
    ]
    for metric, value, status in rows_data:
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(metric).font.size = Pt(9)
        cells[1].paragraphs[0].add_run(value).font.size  = Pt(9)
        r_status = cells[2].paragraphs[0].add_run(status)
        r_status.font.size  = Pt(9)
        r_status.font.color.rgb = RED if status == "⚠" else GREEN

    doc.add_paragraph()

    # ── Issues detected ──────────────────────────────────────────────────
    _heading(doc, "Issues Detected")
    if triggered_rules:
        for rule in triggered_rules:
            _kv(doc, rule.get("name", "Unknown"), rule.get("severity", "medium").upper(),
                value_color=RED if rule.get("severity") == "critical" else BLUE)
    else:
        p = doc.add_paragraph("No hardware threshold violations detected.")
        p.runs[0].font.color.rgb = GREEN

    # ── AI Diagnosis ─────────────────────────────────────────────────────
    if diagnosis:
        _heading(doc, "AI Diagnosis")
        _kv(doc, "Summary",    diagnosis.get("diagnosis", "—"))
        _kv(doc, "Confidence", f"{round(diagnosis.get('confidence', 0) * 100)}%")
        _kv(doc, "Severity",   diagnosis.get("severity", "—").upper())
        _kv(doc, "Engine",     diagnosis.get("llm_model", "—"))

        steps = diagnosis.get("fix_steps", [])
        if steps:
            p = doc.add_paragraph()
            p.add_run("Recommended Fix Steps:").bold = True
            for i, step in enumerate(steps, 1):
                sp = doc.add_paragraph(style="List Number")
                sp.add_run(str(step) if not isinstance(step, dict) else step.get("instruction", str(step)))
                sp.runs[0].font.size = Pt(10)

    # ── Fix execution results ─────────────────────────────────────────────
    if fix_results:
        _heading(doc, "Auto-Fix Execution")
        for fr in fix_results:
            _kv(doc, "Rule",   fr.get("rule_id", "—"))
            _kv(doc, "Method", fr.get("method", "—"))
            _kv(doc, "Result", "SUCCESS ✓" if fr.get("success") else "FAILED ✗",
                value_color=GREEN if fr.get("success") else RED)
            if fr.get("output"):
                code = doc.add_paragraph()
                code.paragraph_format.left_indent = Inches(0.3)
                r = code.add_run(fr["output"][:300])
                r.font.size = Pt(8)
                r.font.name = "Courier New"
                r.font.color.rgb = GREY

    # ── Footer ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("Generated by TED — Tech Express Desk  |  Powered by Vision AI")
    fr.font.size  = Pt(8)
    fr.font.color.rgb = GREY
    fr.italic = True

    # ── Save ─────────────────────────────────────────────────────────────
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename  = f"TED_Incident_{session_id[:8].upper()}_{timestamp}.docx"
    filepath  = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath
